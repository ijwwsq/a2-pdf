"""Сборка .docx в оформлении выбранного бренда из тех же блоков, что и PDF.

Word собирается нативными средствами: обложка, колонтитулы, заголовки, списки,
таблицы, код и цитаты — редактируемый текст, а не картинки. Оформление
сдержанное и деловое: фирменные цвета, логотип, тонкие линии, без фонов
и фотографий (вся визуальная часть живёт в PDF).

Картинками вставляются только диаграммы mermaid — их Word рисовать не умеет.
"""
from __future__ import annotations

import base64
import html as html_mod
import io
import pathlib
import re
import urllib.request

import pymupdf
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

from . import brands, core

WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DPI = 150


def _rgb(color: str) -> RGBColor:
    color = color.lstrip("#")
    return RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))


class Theme:
    """Цвета и шрифты бренда в виде, удобном для python-docx.

    В Word нельзя рассчитывать на фирменные шрифты: если их нет в системе,
    документ подставит засечки. Поэтому берём системные, близкие по духу.
    """

    def __init__(self, brand: brands.Brand,
                 fonts: brands.Fonts | None = None) -> None:
        self.brand = brand
        fonts = fonts or brand.fonts
        self.main = _rgb(brand.color("brand"))
        self.accent = _rgb(brand.color("accent"))
        self.accent_dark = _rgb(brand.color("accent_dark"))
        self.mark = _rgb(brand.color("mark_dark"))
        self.ink = _rgb(brand.color("ink"))
        self.gray = _rgb(brand.neutrals["n500"])
        self.light = _rgb(brand.neutrals["n400"])
        self.hex_main = brand.color("brand").lstrip("#")
        self.hex_mark = brand.color("mark").lstrip("#")
        self.hex_accent = brand.color("accent").lstrip("#")
        self.hex_line = brand.neutrals["n200"].lstrip("#")
        self.hex_soft = brand.neutrals["n50"].lstrip("#")
        self.hex_mark_bg = brand.color("mark_50").lstrip("#")
        self.fonts = fonts
        self.font = fonts.word_body
        self.display = fonts.word_display
        self.mono = fonts.word_mono


# --------------------------------------------------------------------------- #
# Мелкие помощники для XML, которого нет в python-docx
# --------------------------------------------------------------------------- #

def _shade(element, color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    element.append(shd)


def _borders(paragraph, **sides) -> None:
    """Границы абзаца: _borders(p, left=('1FA8FC', 18), bottom=('E1E4EA', 6))."""
    pbdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        if side not in sides:
            continue
        color, size = sides[side]
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), color)
        pbdr.append(node)
    paragraph._p.get_or_add_pPr().append(pbdr)


def _cell_borders(cell, color: str = "E1E4EA", size: int = 4,
                  sides: tuple[str, ...] = ("top", "bottom")) -> None:
    props = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        if side in sides:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), str(size))
            node.set(qn("w:color"), color)
        else:
            node.set(qn("w:val"), "nil")
        borders.append(node)
    props.append(borders)


def _repeat_header(row) -> None:
    """Шапка таблицы повторяется на каждой странице."""
    props = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    props.append(header)


def _field(paragraph, theme: "Theme", code: str, size: float = 8,
           color: RGBColor | None = None):
    """Поле Word — например номер страницы, который считается сам."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    run.font.name = theme.mono
    run.font.size = Pt(size)
    run.font.color.rgb = color or theme.main
    return run


def _spacing(paragraph, before: float = 0, after: float = 0,
             line: float | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line:
        fmt.line_spacing = line


# --------------------------------------------------------------------------- #
# Текст
# --------------------------------------------------------------------------- #

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _style_run(run, size: float, color: RGBColor, name: str) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr().get_or_add_rFonts()
    rpr.set(qn("w:eastAsia"), name)
    rpr.set(qn("w:cs"), name)


def _runs(paragraph, theme: "Theme", text: str, size: float = 10.5,
          color: RGBColor | None = None, bold: bool = False,
          font: str | None = None) -> None:
    """Разбирает **жирный**, *курсив* и `код` в отдельные run-ы."""
    color = color or theme.ink
    font = font or theme.font
    for part in INLINE_RE.split(str(text)):
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
            _style_run(run, size, color, font)
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            _style_run(run, size - 1.5, theme.main, theme.mono)
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
            _style_run(run, size, color, font)
        else:
            run.text = part
            _style_run(run, size, color, font)
        if bold:
            run.bold = True


def _plain(text: str) -> str:
    return re.sub(r"[*`]", "", str(text))


# --------------------------------------------------------------------------- #
# Диаграммы: их рисует браузер, Word получает картинку
# --------------------------------------------------------------------------- #

def _content_box(page) -> pymupdf.Rect:
    """Прямоугольник вокруг нарисованного — чтобы не тащить поля страницы."""
    box = pymupdf.Rect()
    for drawing in page.get_drawings():
        box |= drawing["rect"]
    for word in page.get_text("words"):
        box |= pymupdf.Rect(word[:4])
    if box.is_empty or box.is_infinite:
        return page.rect
    box += (-6, -6, 6, 6)          # немного воздуха вокруг схемы
    return box & page.rect


def _pages_to_png(html_text: str, chrome: str, name: str,
                  wait_ms: int = 15000,
                  crop: bool = False) -> list[tuple[bytes, float]]:
    """Возвращает картинки страниц и соотношение сторон каждой."""
    core.TMP.mkdir(parents=True, exist_ok=True)
    html_path = core.TMP / f"{name}.html"
    pdf_path = core.TMP / f"{name}.pdf"
    html_path.write_text(html_text, encoding="utf-8")
    try:
        core.print_pdf(html_path, pdf_path, chrome, wait_ms)
        images = []
        with pymupdf.open(pdf_path) as doc:
            for page in doc:
                clip = _content_box(page) if crop else None
                pixmap = page.get_pixmap(dpi=DPI, clip=clip)
                images.append((pixmap.tobytes("png"),
                               pixmap.height / max(pixmap.width, 1)))
        return images
    finally:
        html_path.unlink(missing_ok=True)
        pdf_path.unlink(missing_ok=True)


def _diagram_images(sources: list[str], theme: "Theme", chrome: str,
                    name: str) -> list[tuple[bytes, float]]:
    if not sources:
        return []
    brand = theme.brand
    fonts = core.fonts_css_path(theme.fonts.key).read_text(encoding="utf-8")
    lib = (core.ASSETS / "mermaid.min.js").read_text(encoding="utf-8")
    body = "".join(
        '<div class="page"><div class="dg dg-mermaid"><pre class="mermaid">'
        f'{html_mod.escape(src)}</pre></div></div>' for src in sources)
    page = (f'<!doctype html><html lang="ru"><head><meta charset="utf-8">'
            f'<style>{fonts}</style>'
            f'<style>{brands.tokens(brand, theme.fonts)}{core.BODY_CSS}'
            '@page{size:A4;margin:8mm}'
            '.page{break-after:page;display:flex;align-items:center;'
            'justify-content:center;height:281mm}'
            '.dg{border:0;background:none;margin:0;padding:0;width:100%}'
            '.dg-mermaid svg{max-height:275mm}'
            f'</style></head><body>{body}'
            f'<script>{lib}</script>'
            f'<script type="module">'
            f'{core.mermaid_init(brand, theme.fonts)}</script>'
            '</body></html>')
    return _pages_to_png(page, chrome, f"{name}-diagrams", crop=True)


def _load_image(src: str) -> bytes | None:
    if src.startswith("data:"):
        _, _, payload = src.partition(",")
        try:
            return base64.b64decode(payload)
        except Exception:
            return None
    if src.startswith(("http://", "https://")):
        try:
            req = urllib.request.Request(src, headers={"User-Agent": "a2pdf"})
            with urllib.request.urlopen(req, timeout=30) as resp:
                return resp.read(8 * 1024 * 1024)
        except Exception:
            return None
    path = pathlib.Path(src)
    return path.read_bytes() if path.is_file() else None


# --------------------------------------------------------------------------- #
# Стили, обложка и колонтитулы
# --------------------------------------------------------------------------- #

def _setup_styles(document: Document, theme: "Theme") -> None:
    normal = document.styles["Normal"]
    normal.font.name = theme.font
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = theme.ink
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.3
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    rpr.set(qn("w:eastAsia"), theme.font)
    rpr.set(qn("w:cs"), theme.font)

    for name, size, color, before in (
            ("Heading 1", 16, theme.main, 18),
            ("Heading 2", 12.5, theme.main, 14),
            ("Heading 3", 11, theme.ink, 10)):
        style = document.styles[name]
        style.font.name = theme.display
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = theme.font
        style.font.size = Pt(10.5)
        style.font.color.rgb = theme.ink
        style.paragraph_format.space_after = Pt(3)


def _wordmark(paragraph, theme: "Theme", width_cm: float = 3.2) -> None:
    """Логотип картинкой из брендбука — не зависит от шрифтов в системе."""
    logo = core.ASSETS / "logo" / f"{theme.brand.logo}-color.png"
    if logo.is_file():
        paragraph.add_run().add_picture(str(logo), width=Cm(width_cm))
        return
    run = paragraph.add_run(theme.brand.name)
    run.bold = True
    _style_run(run, width_cm * 3.4, theme.main, theme.display)


def _accent_bar(document: Document, theme: "Theme", width_cm: float = 2.2) -> None:
    """Короткий цветной штрих — тот же акцент, что на обложке PDF."""
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(width_cm)
    _shade(cell._tc.get_or_add_tcPr(), theme.hex_mark)
    _cell_borders(cell, sides=())
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    _spacing(paragraph, after=0, line=0.2)
    run = paragraph.add_run(" ")
    _style_run(run, 2, WHITE, theme.font)


def _cover(document: Document, theme: "Theme", front: dict) -> None:
    """Деловая титульная страница: акцент, логотип, заголовок и реквизиты."""
    section = document.sections[0]
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.4)
    section.left_margin = section.right_margin = Cm(2.4)

    _accent_bar(document, theme)

    logo = document.add_paragraph()
    _spacing(logo, before=10, after=4)
    _wordmark(logo, theme, theme.brand.logo_width_cm)

    rule = document.add_paragraph()
    _spacing(rule, before=0, after=0)
    _borders(rule, bottom=(theme.hex_line, 6))

    spacer = document.add_paragraph()
    _spacing(spacer, after=140)

    kicker_parts = []
    if front.get("index"):
        kicker_parts.append(str(front["index"]))
    if front.get("kicker"):
        kicker_parts.append(str(front["kicker"]))
    if kicker_parts:
        line = document.add_paragraph()
        _spacing(line, after=8)
        run = line.add_run("   ".join(kicker_parts).upper())
        _style_run(run, 8.5, theme.accent_dark, theme.mono)

    title = document.add_paragraph()
    _spacing(title, after=8, line=1.05)
    run = title.add_run(_plain(front.get("title", "")))
    run.bold = True
    _style_run(run, 30, theme.main, theme.display)

    if front.get("subtitle"):
        subtitle = document.add_paragraph()
        _spacing(subtitle, after=10, line=1.25)
        _runs(subtitle, theme, str(front["subtitle"]), size=12.5,
              color=theme.gray)

    if front.get("confidential"):
        chip = document.add_paragraph()
        _spacing(chip, before=6, after=0)
        _borders(chip, left=(theme.hex_mark, 18))
        chip.paragraph_format.left_indent = Cm(0.3)
        run = chip.add_run(str(front["confidential"]).upper())
        run.bold = True
        _style_run(run, 8.5, theme.mark, theme.mono)

    meta = list((front.get("meta") or {}).items())
    tail = document.add_paragraph()
    _spacing(tail, after=0)
    tail.paragraph_format.space_before = Pt(120 if not meta else 80)

    if meta:
        table = document.add_table(rows=2, cols=len(meta))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for index, (key, value) in enumerate(meta):
            head_cell = table.rows[0].cells[index]
            head_cell.paragraphs[0].text = ""
            _spacing(head_cell.paragraphs[0], after=2)
            run = head_cell.paragraphs[0].add_run(str(key).upper())
            _style_run(run, 7.5, theme.light, theme.mono)
            _cell_borders(head_cell, color=theme.hex_line, sides=("top",))

            value_cell = table.rows[1].cells[index]
            value_cell.paragraphs[0].text = ""
            _spacing(value_cell.paragraphs[0], after=0)
            run = value_cell.paragraphs[0].add_run(_plain(value))
            _style_run(run, 10, theme.main, theme.mono)
            _cell_borders(value_cell, sides=())

    document.add_paragraph()
    table = document.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    for cell in (left, right):
        _cell_borders(cell, sides=())
        cell.paragraphs[0].text = ""
        _spacing(cell.paragraphs[0], after=0)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    place = left.paragraphs[0].add_run(str(front.get("place", theme.brand.place)))
    _style_run(place, 8, theme.light, theme.mono)
    site = right.paragraphs[0].add_run(theme.brand.site)
    _style_run(site, 8, theme.main, theme.mono)


def _runner_table(container, width_cm: float):
    """Строка колонтитула: слева логотип или подпись, справа — текст."""
    table = container.add_table(rows=1, cols=2, width=Cm(width_cm))
    table.autofit = False
    table.columns[0].width = Cm(width_cm * 0.5)
    table.columns[1].width = Cm(width_cm * 0.5)
    left, right = table.rows[0].cells
    left.width = Cm(width_cm * 0.5)
    right.width = Cm(width_cm * 0.5)
    for cell in (left, right):
        _cell_borders(cell, sides=())
        cell.paragraphs[0].text = ""
        _spacing(cell.paragraphs[0], after=0)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table, left.paragraphs[0], right.paragraphs[0]


def _runners(document: Document, theme: "Theme", front: dict) -> None:
    """Колонтитулы: логотип и название сверху, подпись и номера страниц снизу."""
    section = document.sections[-1]
    section.is_linked_to_previous = False  # титульная остаётся без колонтитулов
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)
    width = Emu(section.page_width - section.left_margin
                - section.right_margin).cm

    header = section.header
    header.paragraphs[0].text = ""
    _spacing(header.paragraphs[0], after=0, line=0.6)
    _, left, right = _runner_table(header, width)
    _wordmark(left, theme, theme.brand.logo_width_cm * 0.55)
    caption = right.add_run(_plain(front.get("header", front.get("title", ""))))
    _style_run(caption, 8, theme.gray, theme.font)
    rule = header.add_paragraph()
    _spacing(rule, before=0, after=0, line=0.6)
    _borders(rule, top=(theme.hex_line, 4))

    footer = section.footer
    footer.paragraphs[0].text = ""
    _spacing(footer.paragraphs[0], after=0, line=0.6)
    _borders(footer.paragraphs[0], bottom=(theme.hex_line, 4))
    _, left, right = _runner_table(footer, width)
    signature = left.add_run(_plain(front.get("footer", "")))
    _style_run(signature, 8, theme.light, theme.font)
    _field(right, theme, "PAGE")
    separator = right.add_run(" / ")
    _style_run(separator, 8, theme.light, theme.mono)
    _field(right, theme, "NUMPAGES", color=theme.light)


# --------------------------------------------------------------------------- #
# Сборка
# --------------------------------------------------------------------------- #

def write_docx(blocks: list[tuple], front: dict, out_path: pathlib.Path,
               chrome: str | None = None, name: str = "document") -> pathlib.Path:
    """Собирает .docx и возвращает путь к файлу."""
    core.ensure_assets(quiet=True)
    chrome = chrome or core.find_chrome()
    front = dict(front)
    brand = brands.get(front.get("brand"))
    theme = Theme(brand, brands.fonts_for(brand, front.get("font")))
    if not front.get("title"):
        first = next((b[1] for b in blocks if b[0] == "h1"), name)
        front["title"] = re.sub(r"^Задание\s+\d+\.\s*", "", str(first))

    document = Document()
    _setup_styles(document, theme)

    with_cover = str(front.get("cover", "true")).lower() not in ("false", "0", "no")
    if with_cover:
        _cover(document, theme, front)
        document.add_section(WD_SECTION.NEW_PAGE)

    section = document.sections[-1]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = section.right_margin = Cm(2.2)
    _runners(document, theme, front)

    diagrams = _diagram_images([b[1] for b in blocks if b[0] == "mermaid"],
                               theme, chrome, name)
    diagram_no = 0
    numbered = str(front.get("numbered", "true")).lower() not in ("false", "0", "no")
    section_no = 0
    content_width = section.page_width - section.left_margin - section.right_margin

    for block in blocks:
        kind = block[0]

        if kind == "h1":
            continue

        if kind == "h2":
            section_no += 1
            paragraph = document.add_paragraph(style="Heading 1")
            _borders(paragraph, top=(theme.hex_line, 4))
            if numbered:
                number = paragraph.add_run(f"{section_no:02d}   ")
                _style_run(number, 10, theme.accent_dark, theme.mono)
            _runs(paragraph, theme, block[1], size=16, color=theme.main,
                  bold=True, font=theme.display)
        elif kind in ("h3", "h4"):
            paragraph = document.add_paragraph(style="Heading 3")
            _runs(paragraph, theme, block[1], size=11, color=theme.ink,
                  bold=True, font=theme.display)
        elif kind == "p":
            _runs(document.add_paragraph(), theme, block[1])
        elif kind in ("ul", "ol"):
            style = "List Bullet" if kind == "ul" else "List Number"
            for item in block[1]:
                _runs(document.add_paragraph(style=style), theme, item)
        elif kind == "code":
            for line in str(block[2]).split("\n"):
                paragraph = document.add_paragraph()
                _spacing(paragraph, after=0, line=1.15)
                paragraph.paragraph_format.left_indent = Cm(0.5)
                _shade(paragraph._p.get_or_add_pPr(), theme.hex_soft)
                _borders(paragraph, left=(theme.hex_accent, 18))
                run = paragraph.add_run(line or " ")
                _style_run(run, 8.5, theme.ink, theme.mono)
            document.add_paragraph()
        elif kind == "note":
            paragraph = document.add_paragraph()
            _spacing(paragraph, before=4, after=8, line=1.25)
            paragraph.paragraph_format.left_indent = Cm(0.5)
            _shade(paragraph._p.get_or_add_pPr(), theme.hex_mark_bg)
            _borders(paragraph, left=(theme.hex_mark, 18))
            _runs(paragraph, theme, block[1], size=10)
        elif kind == "cap":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(paragraph, before=2, after=10)
            run = paragraph.add_run(_plain(block[1]))
            run.italic = True
            _style_run(run, 8.5, theme.gray, theme.font)
        elif kind == "hr":
            paragraph = document.add_paragraph()
            _spacing(paragraph, before=6, after=6)
            _borders(paragraph, bottom=(theme.hex_line, 6))
        elif kind == "table":
            head, rows = block[1], block[2]
            table = document.add_table(rows=1, cols=len(head))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _repeat_header(table.rows[0])
            for cell, title in zip(table.rows[0].cells, head):
                _shade(cell._tc.get_or_add_tcPr(), theme.hex_main)
                _cell_borders(cell, color=theme.hex_main, sides=())
                cell.paragraphs[0].text = ""
                _spacing(cell.paragraphs[0], before=3, after=3)
                run = cell.paragraphs[0].add_run(_plain(title))
                run.bold = True
                _style_run(run, 9, WHITE, theme.font)
            for row in rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row):
                    _cell_borders(cell, color=theme.hex_line, sides=("bottom",))
                    cell.paragraphs[0].text = ""
                    _spacing(cell.paragraphs[0], before=3, after=3, line=1.2)
                    _runs(cell.paragraphs[0], theme, str(value), size=9.5)
            document.add_paragraph()
        elif kind == "mermaid":
            if diagram_no < len(diagrams):
                blob, ratio = diagrams[diagram_no]
                width = content_width
                tallest = Cm(17)          # схема не должна занимать весь лист
                if width * ratio > tallest:
                    width = Emu(int(tallest / ratio))
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(paragraph, before=6, after=4)
                paragraph.add_run().add_picture(io.BytesIO(blob), width=width)
                diagram_no += 1
        elif kind == "image":
            data = _load_image(block[1])
            if data:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(paragraph, before=6, after=6)
                try:
                    paragraph.add_run().add_picture(io.BytesIO(data),
                                                    width=content_width)
                except Exception:
                    pass

    document.core_properties.title = _plain(front["title"])
    document.core_properties.author = theme.brand.name
    document.core_properties.comments = _plain(front.get("subtitle", ""))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)
    return out_path
